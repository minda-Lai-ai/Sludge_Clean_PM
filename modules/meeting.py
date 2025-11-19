import streamlit as st
import pandas as pd
from datetime import date
import io

def render_meeting(project_no):
    st.header(f"{project_no} 專案會議紀錄")
    excel_path = f"Meeting_{project_no}.xlsx"
    columns = ["日期", "地點", "主題", "主持人", "出席人員", "會議記錄"]

    # 讀取現有 Excel
    try:
        df = pd.read_excel(excel_path)
    except Exception:
        df = pd.DataFrame(columns=columns)
    
    # 新增會議
    with st.form("meeting_form"):
        meeting_date = st.date_input("日期", value=date.today())
        location = st.text_input("地點")
        topic = st.text_input("主題")
        host = st.text_input("主持人")
        attendees = st.text_input("出席人員")
        notes = st.text_area("會議記錄")
        submit = st.form_submit_button("新增會議")
        if submit:
            new_row = {
                "日期": meeting_date, "地點": location, "主題": topic, 
                "主持人": host, "出席人員": attendees, "會議記錄": notes
            }
            df = pd.concat([df, pd.DataFrame([new_row])], ignore_index=True)
            df.to_excel(excel_path, index=False)
            st.success("已儲存")

    # 顯示並編輯表格
    if not df.empty:
        st.dataframe(df, use_container_width=True)

        # 修改與刪除
        selected = st.number_input('請輸入要編輯/刪除的行號(0起算)', min_value=0, max_value=len(df)-1, step=1)
        edit = st.button("修改選定列")
        delete = st.button("刪除選定列")
        
        if edit:
            with st.form("edit_form"):
                meeting_date_update = st.date_input("新日期", value=pd.to_datetime(df.loc[selected, "日期"]))
                location_update = st.text_input("新地點", value=df.loc[selected, "地點"])
                topic_update = st.text_input("新主題", value=df.loc[selected, "主題"])
                host_update = st.text_input("新主持人", value=df.loc[selected, "主持人"])
                attendees_update = st.text_input("新出席人員", value=df.loc[selected, "出席人員"])
                notes_update = st.text_area("新會議記錄", value=df.loc[selected, "會議記錄"])
                update = st.form_submit_button("確認修改")
                if update:
                    df.loc[selected] = [meeting_date_update, location_update, topic_update, host_update, attendees_update, notes_update]
                    df.to_excel(excel_path, index=False)
                    st.success("已修改")
        
        if delete:
            df.drop(selected, inplace=True)
            df.reset_index(drop=True, inplace=True)
            df.to_excel(excel_path, index=False)
            st.success("已刪除")

        # 下載: 下載為 Excel
        buffer = io.BytesIO()
        df.to_excel(buffer, index=False)
        buffer.seek(0)
        st.download_button("下載 Excel", data=buffer, file_name=excel_path, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # 下載: 下載為 Word (docx)
        from docx import Document
        doc = Document()
        doc.add_heading(f"{project_no} 會議紀錄", 0)
        table = doc.add_table(rows=1, cols=len(columns))
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(columns):
            hdr_cells[i].text = col_name
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        st.download_button("下載 Word", data=doc_buffer, file_name=f"Meeting_{project_no}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
